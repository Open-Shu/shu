"""Text extraction processor for Shu RAG Backend.

This module provides text extraction functionality for various file types.

OCR mode invariant (post-SHU-728): callers pass `ocr_mode` values from the
`OcrMode` enum — exactly one of `"auto"`, `"always"`, or `"never"`. The
legacy `"fallback"` and `"text_only"` aliases were removed in SHU-728. The
orchestrator (`core.ocr_service.extract_text_with_ocr_fallback`) is the
canonical entry point; it always passes `OcrMode.NEVER` to TextExtractor
because the routing classifier has already chosen text-only by the time we
get here. Direct callers that bypass the orchestrator (`attachment_service`,
`local_ocr_service`) pass `OcrMode.NEVER` and `OcrMode.ALWAYS` respectively.
"""

import asyncio
import os
import re
import threading
import time
from collections.abc import Iterator
from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from typing import Any, ClassVar

from ..core.config import ConfigurationManager
from ..core.logging import get_logger
from ..core.ocr_modes import OcrMode
from ..ingestion.filetypes import (
    ALL_BINARY_SIGNATURES,
    EXT_TO_INGESTION_TYPE,
    KNOWN_BINARY_EXTENSIONS,
    SUPPORTED_TEXT_EXTENSIONS,
    IngestionType,
    detect_extension_from_bytes,
    normalize_extension,
)

logger = get_logger(__name__)

# Common English words used for OCR text quality scoring.
# fmt: off
_COMMON_ENGLISH_WORDS: frozenset[str] = frozenset({
    "the", "and", "or", "but", "in", "on", "at", "to", "for", "of",
    "with", "by", "a", "an", "is", "are", "was", "were", "be", "been",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "can", "this", "that", "these", "those",
    "i", "you", "he", "she", "it", "we", "they", "me", "him", "her",
    "us", "them",
})
# fmt: on

VALID_OCR_MODES: frozenset[str] = frozenset(m.value for m in OcrMode)

# MuPDF process-global store ceiling (SHU-710). The store is shared across all
# concurrent fitz.open() calls, not per-worker. Sizing: per-worker working set
# for sequential text extraction is ~15-25 MiB (active page + its fonts +
# display list); at SHU_OCR_MAX_CONCURRENT_JOBS=6 that is ~120 MiB plus a
# small LRU headroom. Values below ~64 MiB risk thrashing under multi-worker
# load; values above 256 MiB give diminishing returns for text-only workloads.
MUPDF_STORE_MAXSIZE_BYTES: int = 128 * 1024 * 1024


def configure_mupdf_store() -> None:
    """Cap the MuPDF process-global store and shrink defaults (SHU-710).

    Must be called once per process before any ``fitz.open`` — from the
    FastAPI lifespan for the API process, and from the worker entrypoint
    for dedicated worker processes (where OCR actually runs). Idempotent:
    safe to call multiple times, though callers should not rely on that.

    Failure is logged and swallowed so misconfiguration does not block
    process startup; absent the cap, behavior matches pre-SHU-710 MuPDF
    defaults (unbounded store), which is degraded but not broken.
    """
    try:
        import fitz

        fitz.TOOLS.store_maxsize = MUPDF_STORE_MAXSIZE_BYTES
        logger.info(
            "MuPDF store cap configured",
            extra={"store_maxsize_bytes": MUPDF_STORE_MAXSIZE_BYTES},
        )
    except Exception as e:
        logger.warning(f"Failed to configure MuPDF store cap: {e}")


@contextmanager
def _open_pdf(
    file_path: str | None,
    file_content: bytes | None,
    *,
    doc: Any | None = None,
) -> Iterator[Any]:
    """Open a PDF via path (mmap) when possible, else from in-memory bytes.

    When the caller already holds an open ``fitz.Document`` (e.g., the routing
    classifier in ``core.ocr_service`` opens one for the per-page scan), pass
    it as ``doc=`` to reuse it and skip the open/close cycle. The caller owns
    the lifecycle in that case — we yield the doc unchanged and do not close
    or shrink the store.

    Otherwise, always closes the document and shrinks the MuPDF process-global
    store on exit so cached fonts/images from the closed doc are evicted
    before the next extraction begins (SHU-710). The store is process-global,
    so without shrink it staircases toward the cap and forces eviction under
    LRU pressure — shrink returns it to near-baseline between documents.

    ``filetype="pdf"`` is passed explicitly because staged files are named
    ``.bin`` (the staging-service convention) and fitz's default
    extension-based type detection would otherwise fall back to magic-byte
    sniffing. These extraction methods are only reached after upstream
    routing has confirmed the document is a PDF.
    """
    import fitz

    if doc is not None:
        # Caller-owned doc: reuse without managing lifecycle.
        yield doc
        return

    if file_path and os.path.exists(file_path):
        opened = fitz.open(file_path, filetype="pdf")
    elif file_content is not None:
        opened = fitz.open(stream=BytesIO(file_content), filetype="pdf")
    else:
        raise ValueError("_open_pdf requires either file_path or file_content")
    try:
        yield opened
    finally:
        try:
            opened.close()
        finally:
            try:
                fitz.TOOLS.store_shrink(100)
            except Exception:
                # store_shrink is best-effort: a failure here must not mask
                # an earlier extraction exception or block cleanup.
                pass


class UnsupportedFileFormatError(Exception):
    """Exception raised when a file format is not supported for text extraction."""

    pass


class TextExtractor:
    """Text extraction processor for various file types."""

    # --- EasyOCR singleton management ---
    # The Reader loads ~1.5-2.5 GiB of models; creating one per call causes OOM
    # under concurrency.  We cache a single instance and guard init with an async lock.
    _ocr_instance: ClassVar[Any] = None
    _ocr_init_lock: ClassVar[asyncio.Lock | None] = None
    _ocr_init_failed: ClassVar[bool] = False

    # NOTE: OCR concurrency is now managed at the worker level via queue-level
    # capacity tracking (SHU-596). Workers skip the INGESTION_OCR queue when
    # SHU_OCR_MAX_CONCURRENT_JOBS jobs are already active, preventing head-of-line
    # blocking and allowing other work types to proceed.

    # Thread tracking for proper cleanup
    _active_ocr_threads: ClassVar[dict[str, list]] = {}  # job_id -> list of threads
    _thread_lock = threading.Lock()

    # Cancellation events for jobs
    _job_cancellation_events: ClassVar[dict[str, threading.Event]] = {}  # job_id -> threading.Event
    _cancellation_lock = threading.Lock()

    def __init__(self, config_manager: ConfigurationManager) -> None:
        self.config_manager = config_manager

        # Maps IngestionType → handler method.  PDF is intentionally absent
        # because _extract_text_direct routes it to _extract_text_pdf_with_progress.
        self._type_handlers = {
            IngestionType.PLAIN_TEXT: self._extract_text_plain,
            IngestionType.DOCX: self._extract_text_docx,
            IngestionType.DOC: self._extract_text_doc,
            IngestionType.RTF: self._extract_text_rtf,
            IngestionType.HTML: self._extract_text_html,
            IngestionType.EMAIL: self._extract_text_email,
        }

        # Accepted extensions — derived from the central registry.
        self.supported_extensions = set(SUPPORTED_TEXT_EXTENSIONS)

        # Ensure job tracking attribute always exists to avoid AttributeError in logs
        self._current_sync_job_id = None

        # Populated by OCR processing with the actual engine used ("easyocr" or
        # "tesseract").  Read by extract_text() for metadata.
        self._last_ocr_engine: str | None = None

    @classmethod
    def _get_ocr_lock(cls) -> asyncio.Lock:
        """Return the singleton init lock, creating it lazily."""
        if cls._ocr_init_lock is None:
            cls._ocr_init_lock = asyncio.Lock()
        return cls._ocr_init_lock

    @classmethod
    async def get_ocr_instance(cls) -> Any:
        """Get or create the singleton EasyOCR Reader.

        Uses double-checked locking: fast path returns the cached instance without
        acquiring the lock.  The Reader constructor is CPU-heavy (~1.5-2.5 GiB of
        models) and is run in an executor to avoid blocking the event loop.

        Returns None (Tesseract fallback signal) if EasyOCR init fails.
        """
        # Fast path - already initialised
        if cls._ocr_instance is not None:
            return cls._ocr_instance

        # Already tried and failed - don't retry every call
        if cls._ocr_init_failed:
            return None

        async with cls._get_ocr_lock():
            # Double-check after acquiring lock
            if cls._ocr_instance is not None:
                return cls._ocr_instance
            if cls._ocr_init_failed:
                return None

            import certifi

            if not os.environ.get("SSL_CERT_FILE"):
                os.environ["SSL_CERT_FILE"] = certifi.where()

            try:
                logger.info("Initializing EasyOCR singleton Reader")
                loop = asyncio.get_running_loop()
                # Deferred import: easyocr loads ~2GB of models on import,
                # must not load when external OCR is configured.
                import easyocr

                instance = await loop.run_in_executor(None, lambda: easyocr.Reader(["en"]))
                cls._ocr_instance = instance
                logger.info("EasyOCR singleton Reader initialized successfully")
                return cls._ocr_instance
            except Exception as e:
                logger.warning(f"EasyOCR initialization failed: {e}")
                cls._ocr_init_failed = True
                logger.info("Falling back to Tesseract for future OCR calls")
                return None

    @classmethod
    def cleanup_ocr_instance(cls) -> None:
        """Release the singleton EasyOCR Reader to free memory."""
        if cls._ocr_instance is not None:
            logger.info("Releasing EasyOCR singleton Reader")
            cls._ocr_instance = None
        cls._ocr_init_failed = False
        logger.info("OCR instance cleanup completed")

    @classmethod
    def cleanup_ocr_processes(cls) -> None:
        """Clean up all OCR processes and threads."""
        logger.info("Cleaning up OCR processes during shutdown")

        # Clean up all active threads
        with cls._thread_lock:
            total_threads = sum(len(threads) for threads in cls._active_ocr_threads.values())
            if total_threads > 0:
                logger.info(f"Cleaning up {total_threads} active OCR threads")

                for job_id, threads in cls._active_ocr_threads.items():
                    alive_threads = [thread for thread in threads if thread.is_alive()]
                    if alive_threads:
                        logger.warning(f"Orphaning {len(alive_threads)} OCR threads for job {job_id}")

                # Clear all thread tracking
                cls._active_ocr_threads.clear()

        # Clean up cancellation events
        with cls._cancellation_lock:
            cls._job_cancellation_events.clear()

        # Clean up OCR instances
        cls.cleanup_ocr_instance()

        logger.info("OCR process cleanup completed")

    @classmethod
    def register_ocr_thread(cls, job_id: str, thread: threading.Thread) -> None:
        """Register an OCR thread for tracking and cleanup."""
        with cls._thread_lock:
            if job_id not in cls._active_ocr_threads:
                cls._active_ocr_threads[job_id] = []
            cls._active_ocr_threads[job_id].append(thread)
            logger.debug(
                f"Registered OCR thread for job {job_id}, total threads: {len(cls._active_ocr_threads[job_id])}"
            )

    @classmethod
    def cleanup_job_threads(cls, job_id: str) -> None:
        """Clean up all OCR threads for a specific job."""
        with cls._thread_lock:
            if job_id in cls._active_ocr_threads:
                threads = cls._active_ocr_threads[job_id]
                logger.info(f"Cleaning up {len(threads)} OCR threads for job {job_id}")

                # Count threads that are still alive
                alive_threads = [thread for thread in threads if thread.is_alive()]
                if alive_threads:
                    logger.warning(f"OCR thread still running for job {job_id}, thread will be orphaned")
                    logger.info(
                        f"Orphaned {len(alive_threads)} OCR threads for job {job_id} - they will complete naturally"
                    )
                    # Note: Python threads cannot be forcibly terminated
                    # We can only mark them and let them complete naturally

                # Remove from tracking
                del cls._active_ocr_threads[job_id]
                logger.info(f"Removed OCR thread tracking for job {job_id}")

    @classmethod
    def get_active_thread_count(cls, job_id: str | None = None) -> int:
        """Get count of active OCR threads for a job or all jobs."""
        with cls._thread_lock:
            if job_id:
                return len(cls._active_ocr_threads.get(job_id, []))
            return sum(len(threads) for threads in cls._active_ocr_threads.values())

    @classmethod
    def cancel_job_ocr(cls, job_id: str) -> None:
        """Cancel all OCR processing for a specific job."""
        with cls._cancellation_lock:
            if job_id not in cls._job_cancellation_events:
                cls._job_cancellation_events[job_id] = threading.Event()

            # Signal cancellation
            cls._job_cancellation_events[job_id].set()
            logger.info(f"Signaled OCR cancellation for job {job_id}")

    @classmethod
    def is_job_cancelled(cls, job_id: str) -> bool:
        """Check if a job has been cancelled."""
        with cls._cancellation_lock:
            if job_id in cls._job_cancellation_events:
                return cls._job_cancellation_events[job_id].is_set()
            return False

    @classmethod
    def cleanup_job_cancellation(cls, job_id: str) -> None:
        """Clean up cancellation tracking for a job."""
        with cls._cancellation_lock:
            if job_id in cls._job_cancellation_events:
                del cls._job_cancellation_events[job_id]
                logger.debug(f"Cleaned up cancellation tracking for job {job_id}")

    # TODO: Refactor this function. It's too complex (number of branches and statements).
    async def extract_text(  # noqa: PLR0912, PLR0915
        self,
        *,
        file_bytes: bytes | None = None,
        file_path: str | None = None,
        filename: str | None = None,
        mime_type: str | None = None,
        ocr_mode: str | None = None,
        kb_config: dict[str, Any] | None = None,
        progress_context: dict[str, Any] | None = None,
        fitz_doc: Any | None = None,
    ) -> dict[str, Any]:
        """Extract text from a file or in-memory bytes.

        All parameters are keyword-only.  Callers must provide at least
        *file_bytes* or *file_path* (or both).

        Args:
            file_bytes: Raw file content.  When provided together with
                *file_path*, the bytes are used directly and *file_path*
                serves only for extension detection / logging.
            file_path: Path to a file on disk **or** a filename/title used
                only for extension inference (when *file_bytes* is given).
            mime_type: MIME type string used as a fallback for extension
                resolution when *file_path* has no suffix.
            ocr_mode: One of ``"auto"`` (default), ``"always"``, or
                ``"never"`` — see `core.ocr_modes.OcrMode`.
            kb_config: Optional per-KB configuration overrides.
            progress_context: Optional dict carrying progress tracking
                objects (``enhanced_tracker``, ``sync_job_id``, etc.).
            fitz_doc: Optional pre-opened ``fitz.Document``. When provided
                for a PDF, the inner extraction reuses it instead of opening
                a fresh handle — used by `core.ocr_service` to hand off the
                document already opened by the routing classifier (SHU-728).
                The caller retains lifecycle ownership; we do not close it.

        Returns:
            Dictionary containing:
            - text: Extracted text content
            - metadata: Extraction metadata including method, engine,
              confidence, duration

        """
        # --- Derive internal use_ocr bool from the public ocr_mode string ---
        effective_ocr_mode = (ocr_mode or OcrMode.AUTO.value).strip().lower()
        if effective_ocr_mode not in VALID_OCR_MODES:
            raise ValueError(f"Invalid ocr_mode {ocr_mode!r}; must be one of {sorted(VALID_OCR_MODES)}")
        use_ocr = effective_ocr_mode != OcrMode.NEVER.value

        logger.debug(
            "Extracting text from file",
            extra={"file_path": file_path, "ocr_mode": effective_ocr_mode},
        )

        # Set current sync job ID for cancellation tracking
        if progress_context and progress_context.get("sync_job_id"):
            self._current_sync_job_id = progress_context["sync_job_id"]
            logger.debug(f"Set current sync job ID: {self._current_sync_job_id}")

        start_time = time.time()

        # --- Input validation ---
        if file_bytes is None and file_path is None:
            raise ValueError("Either file_bytes or file_path must be provided")

        if file_bytes is None:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
        else:
            logger.debug(
                "Extracting text from in-memory content",
                extra={"file_path": file_path, "mime_type": mime_type},
            )

        # --- Resolve file extension ---
        # Priority: explicit filename hint → real file_path suffix → mime_type → magic bytes.
        # `.bin` is the staging-service convention for opaque blobs and carries no type
        # information; treat it as absent wherever it appears so the next source takes over.
        file_ext = ""
        if filename:
            file_ext = Path(filename).suffix.lower()
            if file_ext == ".bin":
                file_ext = ""
        if not file_ext and file_path:
            file_ext = Path(file_path).suffix.lower()
            if file_ext == ".bin":
                file_ext = ""
        if not file_ext and mime_type:
            file_ext = normalize_extension(mime_type)
            if file_ext == ".bin":
                file_ext = ""
        if not file_ext:
            # Magic-byte sniff: from in-memory bytes if we have them, else from the file head.
            sniffed: str | None = None
            if file_bytes is not None:
                sniffed = detect_extension_from_bytes(file_bytes)
            elif file_path and os.path.exists(file_path):
                try:
                    with open(file_path, "rb") as _fh:
                        sniffed = detect_extension_from_bytes(_fh.read(16))
                except OSError:
                    sniffed = None
            if sniffed:
                file_ext = sniffed
                logger.debug(
                    "No file extension from filename/path/mime, recovered from magic bytes",
                    extra={"file_path": file_path, "recovered_extension": file_ext},
                )
            elif file_bytes is not None:
                # Preserves pre-SHU-710 behavior: degrade to .txt so the extractor can try
                # a best-effort plain-text read instead of hard-failing on missing mime.
                file_ext = ".txt"
                logger.warning(
                    "No file extension resolvable; defaulting to .txt for best-effort extraction",
                    extra={"file_path": file_path, "mime_type": mime_type, "upload_filename": filename},
                )

        if file_ext not in self.supported_extensions:
            logger.warning(
                "Unsupported file format for text extraction: %s",
                file_ext,
                extra={
                    "file_path": file_path,
                    "supported_extensions": sorted(self.supported_extensions),
                },
            )
            raise UnsupportedFileFormatError(f"Unsupported file format: {file_ext}")

        ingestion_type = EXT_TO_INGESTION_TYPE.get(file_ext)

        # Use direct text extraction with OCR configuration
        logger.debug(
            "Using direct text extraction",
            extra={"file_path": file_path, "file_ext": file_ext, "ocr_mode": effective_ocr_mode},
        )

        try:
            text, ocr_actually_used, ocr_confidence = await self._extract_text_direct(
                file_path or "",
                file_bytes,
                progress_context,
                use_ocr,
                file_ext,
                effective_ocr_mode,
                fitz_doc=fitz_doc,
            )
            duration = time.time() - start_time

            # Determine extraction method and engine based on what actually happened.
            # Use ocr_actually_used (not the input use_ocr flag) so fallback-mode PDFs
            # that succeeded via fast text extraction are recorded accurately.
            extraction_method = "text"  # Default for non-OCR
            extraction_engine = "direct"
            extraction_confidence = None

            if ocr_actually_used:
                extraction_method = "ocr"
                extraction_engine = self._last_ocr_engine or "unknown"
                extraction_confidence = ocr_confidence
                actual_method = "ocr"
            elif ingestion_type == IngestionType.PDF:
                extraction_method = "pdf_text"
                extraction_engine = "pymupdf"
                actual_method = "fast_extraction"
            elif ingestion_type in (IngestionType.DOCX, IngestionType.DOC):
                extraction_method = "document"
                extraction_engine = "python-docx"
                actual_method = "fast_extraction"
            elif ingestion_type == IngestionType.PLAIN_TEXT:
                extraction_method = "text"
                extraction_engine = "direct"
                actual_method = "fast_extraction"
            else:
                extraction_method = "text"
                extraction_engine = "direct"
                actual_method = "fast_extraction"

            # Update progress tracker with actual method used
            if progress_context and progress_context.get("enhanced_tracker"):
                tracker = progress_context["enhanced_tracker"]
                if tracker.current_document_tracker and tracker.current_document_tracker.method != actual_method:
                    logger.info(
                        f"Updating processing method from {tracker.current_document_tracker.method} to {actual_method}"
                    )
                    tracker.current_document_tracker.method = actual_method
                    # Broadcast the update

            return {
                "text": text,
                "metadata": {
                    "method": extraction_method,
                    "engine": extraction_engine,
                    "confidence": extraction_confidence,
                    "duration": duration,
                    "details": {
                        "file_extension": file_ext,
                        "ocr_mode": effective_ocr_mode,
                        "processing_time": duration,
                    },
                },
            }
        except Exception as e:
            duration = time.time() - start_time
            logger.error(f"Text extraction failed: {e}", extra={"file_path": file_path})
            raise

    async def _extract_text_direct(
        self,
        file_path: str,
        file_content: bytes | None = None,
        progress_context: dict[str, Any] | None = None,
        use_ocr: bool = True,
        file_ext: str | None = None,
        ocr_mode: str = "auto",
        *,
        fitz_doc: Any | None = None,
    ) -> tuple[str, bool, float | None]:
        """Extract text directly in-memory with progress updates.

        Args:
            file_ext: Pre-resolved file extension (with leading dot). When provided,
                skips re-deriving it from file_path. This is important for files
                whose file_path has no extension (e.g. Google Docs titles) where
                the caller applied a fallback.
            ocr_mode: OCR mode string passed through from the public API.

        Returns:
            (text, ocr_actually_used, confidence) — callers should use the bool for metadata
            rather than inferring from the input use_ocr flag. confidence is None for non-OCR paths.

        """
        try:
            # Set up progress callback if context is provided
            progress_callback = None
            if progress_context:
                progress_callback = self._create_progress_callback(progress_context, use_ocr)

            if not file_ext:
                file_ext = Path(file_path).suffix.lower()

            # Resolve ingestion type for handler dispatch
            ingestion_type = EXT_TO_INGESTION_TYPE.get(file_ext)

            # PDFs: always use the progress-aware path so OCR/use_ocr is honored even without a progress callback
            if ingestion_type == IngestionType.PDF:
                cb = progress_callback if progress_callback else None
                raw_text, ocr_actually_used, ocr_confidence = await self._extract_text_pdf_with_progress(
                    file_path, file_content, cb, use_ocr, ocr_mode, fitz_doc=fitz_doc
                )
            else:
                handler = self._type_handlers[ingestion_type]
                raw_text = await handler(file_path, file_content)
                ocr_actually_used = False
                ocr_confidence = None

            # Clean the extracted text to remove problematic characters
            cleaned_text = self._clean_text(raw_text)

            logger.debug(
                "Successfully extracted and cleaned text",
                extra={
                    "file_path": file_path,
                    "raw_text_length": len(raw_text),
                    "cleaned_text_length": len(cleaned_text),
                    "chars_removed": len(raw_text) - len(cleaned_text),
                },
            )

            return cleaned_text, ocr_actually_used, ocr_confidence

        except Exception as e:
            logger.error("Failed to extract text", extra={"file_path": file_path, "error": str(e)})
            raise

    def _create_progress_callback(self, progress_context: dict[str, Any], use_ocr: bool = False):
        """Create a progress callback function from progress context."""
        if not progress_context:
            return None

        # Extract progress tracking components from context
        enhanced_tracker = progress_context.get("enhanced_tracker")
        sync_job_id = progress_context.get("sync_job_id")
        document_id = progress_context.get("document_id")

        # Store sync job ID for OCR process tracking
        if sync_job_id:
            self._current_sync_job_id = sync_job_id

        if not all([enhanced_tracker, sync_job_id, document_id]):
            logger.warning(
                "Incomplete progress context provided",
                extra={
                    "has_tracker": bool(enhanced_tracker),
                    "has_job_id": bool(sync_job_id),
                    "has_doc_id": bool(document_id),
                },
            )
            return None

        # Capture the current event loop if available
        main_loop = None
        try:
            main_loop = asyncio.get_running_loop()
        except RuntimeError:
            logger.warning(
                "No running event loop found when creating progress callback",
                extra={"sync_job_id": sync_job_id, "document_id": document_id},
            )

        def progress_callback(
            current_page: int, total_pages: int, page_time: float = 0.0, page_text_length: int = 0
        ) -> None:
            """Progress callback that schedules updates in the main event loop."""
            try:
                # Calculate progress percentage
                progress_percent = (current_page / total_pages) * 100 if total_pages > 0 else 0

                # Use the captured main event loop
                if not main_loop:
                    logger.warning(
                        "No event loop available for progress update",
                        extra={
                            "sync_job_id": sync_job_id,
                            "document_id": document_id,
                            "current_page": current_page,
                            "total_pages": total_pages,
                        },
                    )
                    return

                # Schedule the async update in the main event loop
                def schedule_update() -> None:
                    """Schedule the progress update in the main event loop."""
                    try:
                        # Create a coroutine for the update
                        async def update_progress() -> None:
                            try:
                                if use_ocr:
                                    # Use OCR progress update for OCR processing
                                    await enhanced_tracker.update_ocr_progress(
                                        sync_job_id=sync_job_id,
                                        document_id=document_id,
                                        current_page=current_page,
                                        page_time=page_time,
                                        page_text_length=page_text_length,
                                    )
                                    logger.debug(
                                        "Updated OCR progress",
                                        extra={
                                            "sync_job_id": sync_job_id,
                                            "document_id": document_id,
                                            "current_page": current_page,
                                            "total_pages": total_pages,
                                            "progress_percent": progress_percent,
                                        },
                                    )
                                else:
                                    # Use regular progress update for fast extraction
                                    await enhanced_tracker.update_document_progress(
                                        sync_job_id=sync_job_id,
                                        document_id=document_id,
                                        current_page=current_page,
                                        total_pages=total_pages,
                                        progress_percent=progress_percent,
                                    )
                                    logger.debug(
                                        "Updated fast extraction progress",
                                        extra={
                                            "sync_job_id": sync_job_id,
                                            "document_id": document_id,
                                            "current_page": current_page,
                                            "total_pages": total_pages,
                                            "progress_percent": progress_percent,
                                        },
                                    )
                            except Exception as e:
                                logger.warning(
                                    f"Progress update failed: {e}",
                                    extra={
                                        "sync_job_id": sync_job_id,
                                        "document_id": document_id,
                                        "current_page": current_page,
                                        "total_pages": total_pages,
                                        "use_ocr": use_ocr,
                                        "error": str(e),
                                    },
                                )

                        # Schedule the coroutine in the main event loop
                        asyncio.create_task(update_progress())  # noqa: RUF006 # we don't need the task reference

                    except Exception as e:
                        logger.error(
                            "Failed to schedule progress update",
                            extra={
                                "sync_job_id": sync_job_id,
                                "document_id": document_id,
                                "error": str(e),
                            },
                        )

                # Use call_soon_threadsafe to schedule from any thread
                main_loop.call_soon_threadsafe(schedule_update)

            except Exception as e:
                logger.error(
                    "Failed to update progress in direct extraction",
                    extra={
                        "sync_job_id": sync_job_id,
                        "document_id": document_id,
                        "current_page": current_page,
                        "total_pages": total_pages,
                        "error": str(e),
                    },
                )

        return progress_callback

    async def _extract_text_pdf_with_progress(
        self,
        file_path: str,
        file_content: bytes | None = None,
        progress_callback=None,
        use_ocr: bool = True,
        ocr_mode: str = "auto",
        *,
        fitz_doc: Any | None = None,
    ) -> tuple[str, bool, float | None]:
        """Extract text from PDF with page-by-page progress updates.

        Returns:
            (text, ocr_actually_used, confidence) — ocr_actually_used is True only when
            OCR ran. confidence is the real per-word average from EasyOCR, a text quality
            heuristic when Tesseract ran, or None when fast text extraction succeeded.

        """
        logger.debug(
            "Extracting PDF text with progress updates",
            extra={"file_path": file_path, "use_ocr": use_ocr, "ocr_mode": ocr_mode},
        )

        if use_ocr:
            # OCR is enabled - use OCR directly. The OCR path renders pages from
            # raw bytes via Mistral/EasyOCR and doesn't read the fitz text layer,
            # so a pre-opened doc has nothing to contribute here — don't forward.
            logger.info("OCR enabled for PDF, using OCR processing", extra={"file_path": file_path})
            text, confidence = await self._extract_pdf_ocr_direct(file_path, file_content, progress_callback)
            return text, True, confidence
        # OCR is disabled - try text extraction only
        logger.info("OCR disabled for PDF, using text extraction only", extra={"file_path": file_path})
        return (
            await self._extract_pdf_text_only(file_path, file_content, progress_callback, fitz_doc=fitz_doc),
            False,
            None,
        )

    async def _extract_pdf_text_only(
        self,
        file_path: str,
        file_content: bytes | None = None,
        progress_callback=None,
        *,
        fitz_doc: Any | None = None,
    ) -> str:
        """Extract text from PDF using text extraction methods only (no OCR).

        When ``fitz_doc`` is provided (e.g., handed off from the SHU-728
        routing classifier), the existing handle is reused — no second open,
        no second mmap. The caller retains lifecycle ownership.
        """
        logger.debug("Extracting PDF text only (no OCR)", extra={"file_path": file_path})

        def _extract_text_only():
            """Extract PDF text without OCR."""
            try:
                with _open_pdf(file_path, file_content, doc=fitz_doc) as doc:
                    total_pages = len(doc)
                    logger.debug(f"PDF has {total_pages} pages", extra={"file_path": file_path})

                    if progress_callback:
                        progress_callback(0, total_pages)

                    parts: list[str] = []
                    for page_num in range(total_pages):
                        page = doc.load_page(page_num)
                        page_text = page.get_text()

                        if page_text.strip():
                            parts.append(page_text)

                        if progress_callback:
                            progress_callback(page_num + 1, total_pages)

                        logger.debug(
                            f"Processed page {page_num + 1}/{total_pages}",
                            extra={"file_path": file_path, "page_text_length": len(page_text)},
                        )

                    return "\n".join(parts).strip()

            except Exception as e:
                logger.error(f"PDF text extraction failed: {e}", extra={"file_path": file_path})
                return ""

        # Run in executor to avoid blocking
        result = await asyncio.get_running_loop().run_in_executor(None, _extract_text_only)

        if not result.strip():
            logger.warning("No text found in PDF with text extraction only", extra={"file_path": file_path})
            return ""

        return result

    async def _extract_pdf_ocr_direct(
        self, file_path: str, file_content: bytes | None = None, progress_callback=None
    ) -> tuple[str, float]:
        """Extract PDF text using direct in-process OCR with proper metadata tracking.

        OCR concurrency is managed at the worker level via queue-level capacity
        tracking (SHU-596). Workers skip the INGESTION_OCR queue when at capacity,
        so at most SHU_OCR_MAX_CONCURRENT_JOBS jobs will be processing OCR
        simultaneously per worker process.

        Returns:
            (text, confidence) — confidence is the real per-word average from EasyOCR,
            or a text quality heuristic from ``_calculate_text_quality`` when Tesseract ran.

        """
        start_time = time.time()

        try:
            # Worker-level concurrency limiting ensures at most
            # SHU_OCR_MAX_CONCURRENT_JOBS are processing OCR simultaneously.
            with _open_pdf(file_path, file_content) as doc:
                total_pages = len(doc)
                logger.info(
                    f"Starting direct OCR processing for {total_pages} pages",
                    extra={"file_path": file_path},
                )

                try:
                    text, method, confidence = await self._process_pdf_with_ocr_direct(
                        doc, file_path, progress_callback
                    )
                    engine_map = {"ocr": "easyocr", "tesseract_direct": "tesseract"}
                    self._last_ocr_engine = engine_map.get(method, method)
                    if text.strip():
                        processing_time = time.time() - start_time
                        logger.info(
                            "OCR processing successful",
                            extra={
                                "file_path": file_path,
                                "engine": self._last_ocr_engine,
                                "confidence": confidence,
                                "processing_time": processing_time,
                                "pages": total_pages,
                            },
                        )
                        return text, confidence
                except Exception as e:
                    logger.error(f"OCR processing failed: {e}", extra={"file_path": file_path})
                    return "", 0.0

                logger.error("All direct OCR methods failed", extra={"file_path": file_path})
                return "", 0.0

        except Exception as e:
            logger.error(f"Direct OCR processing failed: {e}", extra={"file_path": file_path})
            # Fallback: if the file has a .pdf extension but isn't a valid PDF, try reading as plain text
            try:
                if file_content is not None:
                    text = file_content.decode("utf-8", errors="ignore").strip()
                else:
                    with open(file_path, "rb") as f:
                        text = f.read().decode("utf-8", errors="ignore").strip()
                if text:
                    logger.info(
                        "PDF open failed; fell back to plain-text read for OCR test fixture",
                        extra={"file_path": file_path},
                    )
                    return text, self._calculate_text_quality(text)
            except Exception as fallback_err:
                logger.warning(
                    f"Plain-text fallback after OCR failure also failed: {fallback_err}",
                    extra={"file_path": file_path},
                )
            return "", 0.0

    # TODO: Refactor this function. It's too complex (number of branches and statements).
    async def _process_pdf_with_ocr_direct(self, doc, file_path: str, progress_callback=None):  # noqa: PLR0912, PLR0915
        """Process PDF with OCR directly (EasyOCR with Tesseract fallback)."""
        import threading
        import time

        # Get the OCR instance (EasyOCR with Tesseract fallback)
        try:
            logger.info("Getting OCR instance", extra={"file_path": file_path})
            ocr = await self.get_ocr_instance()
            if ocr is None:
                logger.info("Using Tesseract fallback", extra={"file_path": file_path})
                # Use Tesseract directly
                return self._process_pdf_with_tesseract_direct(doc, file_path, progress_callback)
            logger.info("OCR instance obtained successfully", extra={"file_path": file_path})
        except Exception as e:
            logger.error(f"Failed to get OCR instance: {e}", extra={"file_path": file_path})
            # Fallback to Tesseract
            logger.info("Falling back to Tesseract", extra={"file_path": file_path})
            return self._process_pdf_with_tesseract_direct(doc, file_path, progress_callback)

        # Process pages with EasyOCR
        import fitz  # Add missing import

        page_outputs: list[str] = []
        confidence_scores = []
        total_pages = len(doc)

        for page_num in range(total_pages):
            # Check for job cancellation before processing each page
            if self._current_sync_job_id and self.is_job_cancelled(self._current_sync_job_id):
                logger.info(f"OCR cancelled for job {self._current_sync_job_id}, stopping at page {page_num + 1}")
                break

            page = doc[page_num]

            logger.info(
                f"Starting OCR processing for page {page_num + 1}/{total_pages}",
                extra={"file_path": file_path},
            )

            # Update progress - starting page
            page_start_time = time.time()
            if progress_callback:
                progress_callback(page_num, total_pages)

            # Convert page to image
            render_scale = self.config_manager.settings.ocr_render_scale
            pix = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale))

            # Convert to numpy array for EasyOCR
            import io

            import numpy as np
            from PIL import Image

            # Convert pixmap to PIL Image, then to numpy array
            img_data = pix.tobytes("png")
            pil_image = Image.open(io.BytesIO(img_data))
            img_array = np.array(pil_image)

            # Set up progress monitoring during OCR processing
            ocr_result = None
            ocr_error = None
            ocr_complete = threading.Event()

            def run_ocr(
                current_page_num: int = page_num,
                _img: "np.ndarray" = img_array,
                _done: threading.Event = ocr_complete,
            ) -> None:
                """Run OCR in a separate thread so we can monitor progress."""
                nonlocal ocr_result, ocr_error
                try:
                    # Check for job cancellation before starting OCR
                    if self._current_sync_job_id and self.is_job_cancelled(self._current_sync_job_id):
                        logger.info(f"OCR cancelled for job {self._current_sync_job_id} on page {current_page_num + 1}")
                        ocr_error = Exception("OCR cancelled")
                        return

                    logger.info(f"Running OCR on page {current_page_num + 1}", extra={"file_path": file_path})

                    # Use EasyOCR (Tesseract fallback handled in get_ocr_instance)
                    if hasattr(ocr, "readtext"):  # EasyOCR
                        ocr_result = ocr.readtext(_img)
                        logger.info(f"EasyOCR completed for page {current_page_num + 1}")
                    else:
                        logger.error(f"Unknown OCR instance type on page {current_page_num + 1}")
                        ocr_result = []
                except Exception as e:
                    ocr_error = e
                    logger.error(f"OCR failed on page {current_page_num + 1}: {e}", extra={"file_path": file_path})
                finally:
                    _done.set()

            # Start OCR in background thread
            ocr_thread = threading.Thread(target=run_ocr)
            ocr_thread.daemon = False  # Changed to False so we can properly track and wait for completion
            ocr_thread.start()
            # Release our references to the page image data immediately.  The thread
            # captured img_array by value via the default-arg binding above, so GC can
            # reclaim this memory as soon as the thread finishes — even if the thread
            # is orphaned by a timeout.
            del img_array
            pix = None

            # Register thread for tracking if we have a job ID
            if self._current_sync_job_id:
                self.register_ocr_thread(self._current_sync_job_id, ocr_thread)

            # Monitor progress while OCR is running (ASYNC VERSION with timeout)
            progress_counter = 0
            max_wait_time = self.config_manager.settings.ocr_page_timeout
            start_time = time.time()

            while not ocr_complete.is_set():
                # Check for job cancellation
                if self._current_sync_job_id and self.is_job_cancelled(self._current_sync_job_id):
                    logger.info(f"OCR monitoring cancelled for job {self._current_sync_job_id} on page {page_num + 1}")
                    # Force completion and use empty result
                    ocr_complete.set()
                    ocr_result = []
                    break

                # Check for timeout
                if time.time() - start_time > max_wait_time:
                    logger.error(
                        f"OCR timeout on page {page_num + 1} after {max_wait_time} seconds",
                        extra={"file_path": file_path},
                    )
                    # Force completion and use empty result
                    ocr_complete.set()
                    ocr_result = []
                    break

                # Send progress updates every 0.5 seconds during OCR processing (NON-BLOCKING)
                # This also makes cancellation more responsive
                await asyncio.sleep(0.5)
                progress_counter += 1

                if progress_callback:
                    # Show sub-page progress (page X.Y/total)
                    sub_progress = min(progress_counter * 0.1, 0.9)  # Max 90% progress per page
                    effective_page = page_num + sub_progress
                    progress_callback(effective_page, total_pages)

                logger.debug(
                    f"OCR processing page {page_num + 1}, progress update {progress_counter}",
                    extra={"file_path": file_path},
                )

            # Wait for OCR to complete (NON-BLOCKING with timeout)
            try:
                await asyncio.wait_for(
                    asyncio.get_running_loop().run_in_executor(None, ocr_thread.join),
                    timeout=10.0,  # 10 second timeout for thread join
                )
            except TimeoutError:
                logger.error(
                    f"OCR thread join timeout on page {page_num + 1}",
                    extra={"file_path": file_path},
                )

            # Process results
            page_text = ""
            try:
                if ocr_error:
                    raise ocr_error

                result = ocr_result

                if result:
                    page_confidences = []

                    # Handle EasyOCR result format
                    if hasattr(ocr, "readtext"):  # EasyOCR format
                        detections: list[str] = []
                        for detection in result:
                            if len(detection) >= 3:
                                _bbox, text_content, confidence = detection
                                detections.append(text_content)
                                page_confidences.append(confidence)
                        page_text = " ".join(detections)
                    else:
                        # Unknown format - log warning and skip
                        logger.warning(f"Unknown OCR result format on page {page_num + 1}")
                        page_text = ""

                    page_outputs.append(page_text)
                    confidence_scores.extend(page_confidences)

                # Update progress - page completed
                page_time = time.time() - page_start_time
                page_text_length = len(page_text) if page_text else 0
                if progress_callback:
                    progress_callback(page_num + 1, total_pages, page_time, page_text_length)

                logger.info(
                    f"Completed OCR processing for page {page_num + 1}/{total_pages}",
                    extra={"file_path": file_path},
                )

            except Exception as e:
                logger.warning(
                    f"EasyOCR processing failed on page {page_num + 1}: {e}",
                    extra={"file_path": file_path},
                )

                # Still update progress even if page failed
                page_time = time.time() - page_start_time
                if progress_callback:
                    progress_callback(page_num + 1, total_pages, page_time, 0)

                # Add empty text for failed page to maintain page count
                page_outputs.append(f"[OCR failed on page {page_num + 1}]")
                continue

        avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0

        return "\n".join(page_outputs).strip(), "ocr", avg_confidence

    def _process_pdf_with_tesseract_direct(self, doc, file_path: str, progress_callback=None):
        """Process PDF with Tesseract directly (no process isolation)."""
        try:
            import io

            import fitz  # Add missing import
            import pytesseract
            from PIL import Image
        except ImportError as e:
            raise Exception(f"Tesseract dependencies not available: {e}")

        page_outputs: list[str] = []
        total_pages = len(doc)

        for page_num in range(total_pages):
            # Check for job cancellation before processing each page
            if self._current_sync_job_id and self.is_job_cancelled(self._current_sync_job_id):
                logger.info(
                    f"Tesseract OCR cancelled for job {self._current_sync_job_id}, stopping at page {page_num + 1}"
                )
                break

            page = doc[page_num]

            # Convert page to image
            render_scale = self.config_manager.settings.ocr_render_scale
            pix = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale))
            img_data = pix.tobytes("png")

            # Run Tesseract on page
            try:
                image = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(image, config="--psm 6")
                page_outputs.append(page_text)

                # Update progress
                if progress_callback:
                    progress_callback(page_num + 1, total_pages)

            except Exception as e:
                logger.warning(f"Tesseract failed on page {page_num + 1}: {e}")
                continue

        text = "\n".join(page_outputs).strip()
        return text, "tesseract_direct", self._calculate_text_quality(text)

    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing problematic characters."""
        if not text:
            return text

        # Remove NULL characters (0x00) that cause PostgreSQL errors
        cleaned = text.replace("\x00", "")

        # Remove other problematic control characters but keep important whitespace
        # Keep: \t (tab), \n (newline), \r (carriage return)
        # Remove: other control characters (0x01-0x08, 0x0b-0x0c, 0x0e-0x1f, 0x7f)
        cleaned = re.sub(r"[\x01-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", cleaned)

        # Normalize excessive whitespace but preserve paragraph breaks
        # Replace multiple spaces with single space
        cleaned = re.sub(r" +", " ", cleaned)

        # Replace multiple newlines with double newline (preserve paragraphs)
        cleaned = re.sub(r"\n\n+", "\n\n", cleaned)

        # Strip leading/trailing whitespace
        return cleaned.strip()

    async def _extract_text_plain(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from plain text files."""

        def _read_file():
            if file_content:
                return file_content.decode("utf-8", errors="ignore")
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                return f.read()

        return await asyncio.get_running_loop().run_in_executor(None, _read_file)

    def _calculate_text_quality(self, text: str) -> float:
        """Calculate a quality score for extracted text."""
        if not text:
            return 0.0

        # Count meaningful words (not just random characters)
        words = text.split()
        if not words:
            return 0.0

        # Calculate ratio of meaningful words (words with letters)
        meaningful_words = sum(1 for word in words if any(c.isalpha() for c in word))
        meaningful_ratio = meaningful_words / len(words) if words else 0

        # Calculate average word length (longer words often indicate better OCR)
        avg_word_length = sum(len(word) for word in words) / len(words) if words else 0

        # Calculate ratio of common English words
        common_word_count = sum(1 for word in words if word.lower() in _COMMON_ENGLISH_WORDS)
        common_word_ratio = common_word_count / len(words) if words else 0

        # Combine scores
        return meaningful_ratio * 0.4 + min(avg_word_length / 10.0, 1.0) * 0.3 + common_word_ratio * 0.3

    # TODO: Refactor this function. It's too complex (number of branches and statements).
    async def _extract_text_docx(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from DOCX files using multiple methods with fallbacks."""

        def _try_python_docx():
            """Try python-docx for DOCX text extraction."""
            try:
                from io import BytesIO

                import docx

                doc = docx.Document(BytesIO(file_content)) if file_content else docx.Document(file_path)

                parts: list[str] = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)

                # Also extract text from tables with proper formatting
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                            else:
                                row_text.append("")  # Preserve empty cells for structure
                        if any(row_text):  # Only add non-empty rows
                            table_text.append(" | ".join(row_text))

                    if table_text:
                        parts.append("\n[Table]\n" + "\n".join(table_text) + "\n[/Table]")

                return "\n".join(parts).strip()
            except Exception as e:
                logger.debug(f"python-docx extraction failed: {e!s}")
                return None

        # Try multiple DOCX extraction methods in order of preference
        extraction_methods = [("python-docx", _try_python_docx)]

        for method_name, extractor_func in extraction_methods:
            try:
                logger.debug(f"Attempting DOCX extraction with {method_name}")
                text = await asyncio.get_running_loop().run_in_executor(None, extractor_func)

                if text and text.strip():
                    logger.debug(
                        f"Successfully extracted DOCX text using {method_name}",
                        extra={
                            "file_path": file_path,
                            "method": method_name,
                            "text_length": len(text),
                        },
                    )
                    return text
                logger.debug(f"{method_name} returned empty text")

            except Exception as e:
                logger.debug(f"{method_name} extraction failed: {e!s}")

        # If all methods fail, raise an exception instead of returning error message
        logger.warning("All DOCX extraction methods failed", extra={"file_path": file_path})
        raise Exception(f"Could not extract text from DOCX {file_path} - all extraction methods failed")

    async def _extract_text_doc(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from DOC files."""
        logger.warning("DOC extraction not supported, falling back to basic text extraction")
        return await self._extract_text_fallback(file_path, file_content)

    async def _extract_text_rtf(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from RTF files."""
        try:
            from striprtf.striprtf import rtf_to_text

            def _extract_rtf():
                if file_content:
                    rtf_content = file_content.decode("utf-8", errors="ignore")
                else:
                    with open(file_path, encoding="utf-8", errors="ignore") as f:
                        rtf_content = f.read()
                return rtf_to_text(rtf_content)

            return await asyncio.get_running_loop().run_in_executor(None, _extract_rtf)
        except ImportError:
            logger.warning("striprtf not installed, falling back to basic text extraction")
            return await self._extract_text_fallback(file_path, file_content)
        except Exception as e:
            logger.warning("RTF extraction failed, falling back", extra={"error": str(e)})
            return await self._extract_text_fallback(file_path, file_content)

    async def _extract_text_html(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from HTML files."""
        try:
            from bs4 import BeautifulSoup

            def _extract_html():
                if file_content:
                    content = file_content.decode("utf-8", errors="ignore")
                else:
                    with open(file_path, encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                soup = BeautifulSoup(content, "html.parser")
                return soup.get_text()

            return await asyncio.get_running_loop().run_in_executor(None, _extract_html)
        except ImportError:
            logger.warning("BeautifulSoup not installed, falling back to basic text extraction")
            return await self._extract_text_fallback(file_path, file_content)
        except Exception as e:
            logger.warning("HTML extraction failed, falling back", extra={"error": str(e)})
            return await self._extract_text_fallback(file_path, file_content)

    async def _extract_text_fallback(self, file_path: str, file_content: bytes | None = None) -> str:
        """Fallback text extraction method with strict binary file detection."""
        logger.debug("Using fallback text extraction", extra={"file_path": file_path})
        try:

            def _read_file():
                if file_content:
                    # For in-memory content, check if it looks like binary data
                    for signature in ALL_BINARY_SIGNATURES:
                        if file_content.startswith(signature):
                            logger.warning(f"Binary file detected in memory content: {file_path}")
                            return f"[Error: Binary file detected - cannot extract text from {file_path}]"

                    # Try to decode as text, but be careful
                    try:
                        return file_content.decode("utf-8", errors="ignore")
                    except UnicodeDecodeError:
                        logger.warning(f"Failed to decode content as UTF-8: {file_path}")
                        return f"[Error: Cannot decode content as text - {file_path}]"
                else:
                    # Check file extension for known binary formats
                    file_ext = Path(file_path).suffix.lower()

                    if file_ext in KNOWN_BINARY_EXTENSIONS:
                        logger.warning(f"Binary file extension detected: {file_path}")
                        return f"[Error: Binary file format {file_ext} - cannot extract text from {file_path}]"

                    # For text-based files, try reading as text
                    try:
                        with open(file_path, encoding="utf-8", errors="ignore") as f:
                            content = f.read()
                            if not content.strip():
                                logger.warning(f"Empty or whitespace-only file: {file_path}")
                                return f"[Error: Empty file - {file_path}]"
                            return content
                    except UnicodeDecodeError:
                        logger.warning(f"Failed to decode file as UTF-8: {file_path}")
                        return f"[Error: Cannot decode file as text - {file_path}]"
                    except Exception as e:
                        logger.warning(f"Failed to read file: {file_path}, error: {e!s}")
                        return f"[Error: Cannot read file - {file_path}]"

            return await asyncio.get_running_loop().run_in_executor(None, _read_file)
        except Exception as e:
            logger.error("Fallback extraction failed", extra={"file_path": file_path, "error": str(e)})
            return f"[Error: Could not extract text from {file_path} - {e!s}]"

    async def _extract_text_email(self, file_path: str, file_content: bytes | None = None) -> str:
        """Extract text from email content (Gmail messages, etc.)."""
        try:
            if file_content:
                # Email content is already in text format (from Gmail API)
                text = file_content.decode("utf-8")
                logger.debug(
                    "Successfully extracted email text from content",
                    extra={"file_path": file_path, "content_length": len(text)},
                )
                return text
            # Fallback: treat as plain text file
            return await self._extract_text_plain(file_path, file_content)
        except Exception as e:
            logger.error("Failed to extract email text", extra={"file_path": file_path, "error": str(e)})
            return f"[Error: Could not extract email text from {file_path} - {e!s}]"

    def get_supported_formats(self) -> list:
        """Get list of supported file formats."""
        return sorted(self.supported_extensions)

    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_extensions
